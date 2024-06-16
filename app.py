import re
from collections import Counter
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

def read_sql_file(filepath):
    with open(filepath, 'r') as file:
        return file.read()

def tokenize_sql(sql_content):
    sql_content = re.sub(r'--.*?(\n|$)|/\*.*?\*/', ' ', sql_content, flags=re.S)
    words = re.findall(r'\b\w+\b', sql_content.lower())
    return words

sql_content1 = read_sql_file('data/pyprc_cmn_pay_rule_proration_method.sql')
sql_content2 = read_sql_file('data/pyprc_hcl_aws_cmn_pay_rule_emp_fetch.sql')

words1 = tokenize_sql(sql_content1)
words2 = tokenize_sql(sql_content2)

combined_words = words1 + words2
word_counts = Counter(combined_words)

filtered_word_counts = {word: count for word, count in word_counts.items() if count >= 2}

most_common_20 = dict(word_counts.most_common(20))
least_common_20 = dict(word_counts.most_common()[:-21:-1])

def plot_word_counts(word_counts, title, filename):
    plt.figure(figsize=(12, 6))
    plt.bar(word_counts.keys(), word_counts.values(), color='skyblue')
    plt.xticks(rotation=90)
    plt.title(title)
    plt.xlabel('Words')
    plt.ylabel('Counts')
    plt.savefig(filename)
    plt.close()

plot_word_counts(most_common_20, '20 Most Occurring Words', 'most_common_20.png')
plot_word_counts(least_common_20, '20 Least Occurring Words', 'least_common_20.png')

def extract_comments(sql_content):
    comments = re.findall(r'--.*?$|/\*.*?\*/', sql_content, re.S | re.M)
    return comments

def extract_variables_and_returns(sql_content):
    variables = re.findall(r'@\w+', sql_content)
    return_values = re.findall(r'RETURN\s+@\w+', sql_content, re.I)
    return variables, return_values

def extract_table_names(sql_content):
    table_names = re.findall(r'\bFROM\s+(\w+)|\bJOIN\s+(\w+)', sql_content, re.I)
    table_names = [name for name_pair in table_names for name in name_pair if name]
    return list(set(table_names))

comments1 = extract_comments(sql_content1)
comments2 = extract_comments(sql_content2)

variables1, return_values1 = extract_variables_and_returns(sql_content1)
variables2, return_values2 = extract_variables_and_returns(sql_content2)

table_names1 = extract_table_names(sql_content1)
table_names2 = extract_table_names(sql_content2)

comments = comments1 + comments2
variables = list(set(variables1 + variables2))
return_values = list(set(return_values1 + return_values2))
table_names = list(set(table_names1 + table_names2))

def extract_mathematical_formulae(sql_content):
    formulae = re.findall(r'[\w@]+\s*[\+\-\*/%]\s*[\w@]+', sql_content)
    return formulae

formulae1 = extract_mathematical_formulae(sql_content1)
formulae2 = extract_mathematical_formulae(sql_content2)

formulae = list(set(formulae1 + formulae2))

def find_max_nesting_level(sql_content):
    nesting_keywords = ['SELECT', 'INSERT', 'UPDATE', 'DELETE', 'IF', 'WHILE', 'CASE', 'BEGIN', 'END']
    tokens = re.findall(r'\b(?:' + '|'.join(nesting_keywords) + r')\b', sql_content, re.I)
    
    max_depth = 0
    current_depth = 0

    for token in tokens:
        if token.upper() in ['SELECT', 'INSERT', 'UPDATE', 'DELETE', 'IF', 'WHILE', 'CASE', 'BEGIN']:
            current_depth += 1
            if current_depth > max_depth:
                max_depth = current_depth
        elif token.upper() == 'END':
            current_depth -= 1
    
    return max_depth

max_nesting_level1 = find_max_nesting_level(sql_content1)
max_nesting_level2 = find_max_nesting_level(sql_content2)

nesting_levels = {
    'SP1': max_nesting_level1,
    'SP2': max_nesting_level2
}

plt.figure(figsize=(8, 4))
plt.bar(nesting_levels.keys(), nesting_levels.values(), color='green')
plt.title('Maximum Nesting Level of SQL Statements')
plt.xlabel('Stored Procedures')
plt.ylabel('Nesting Level')
plt.savefig('nesting_levels.png')
plt.close()

doc = Document()
doc.add_heading('SQL Analysis Results', level=1)

doc.add_heading('20 Most Occurring Words', level=2)
doc.add_picture('most_common_20.png', width=Inches(6))

doc.add_heading('20 Least Occurring Words', level=2)
doc.add_picture('least_common_20.png', width=Inches(6))

doc.add_heading('Comments', level=2)
for comment in comments:
    doc.add_paragraph(comment)

doc.add_heading('Variables', level=2)
for variable in variables:
    doc.add_paragraph(variable)

doc.add_heading('Return Values', level=2)
for return_value in return_values:
    doc.add_paragraph(return_value)

doc.add_heading('Table Names', level=2)
for table_name in table_names:
    doc.add_paragraph(table_name)

doc.add_heading('Mathematical Formulae', level=2)
for formula in formulae:
    doc.add_paragraph(formula)

doc.add_heading('Maximum Nesting Level', level=2)
doc.add_picture('nesting_levels.png', width=Inches(6))

doc.save('SQL_Analysis_Results.docx')
